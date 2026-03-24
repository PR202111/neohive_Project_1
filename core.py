from __future__ import annotations

import json
import os
import re
import uuid
from pathlib import Path
from typing import List, Optional

import requests
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches
from pydantic import BaseModel, Field, field_validator
import logging
from load_dotenv import load_dotenv


load_dotenv()

# ---------------------------------------------------------------------------
# Logging
# ---------------------------------------------------------------------------
logging.basicConfig(level=logging.INFO, format="%(levelname)s | %(message)s")
logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------
OPENROUTER_API_KEY: str = os.getenv("OPENROUTER_API_KEY", "")
OPENROUTER_BASE_URL: str = "https://openrouter.ai/api/v1/chat/completions"
DEFAULT_MODEL = "meta-llama/llama-3-8b-instruct"
OUTPUT_DIR: Path = Path("outputs")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


# ---------------------------------------------------------------------------
# Pydantic Models
# ---------------------------------------------------------------------------

class Project(BaseModel):
    name: Optional[str] = None
    description: Optional[List[str]] = Field(default_factory=list)
    technologies: Optional[List[str]] = Field(default_factory=list)
    link: Optional[str] = None

    @field_validator("description", "technologies", mode="before")
    @classmethod
    def coerce_list(cls, v):
        if v is None:
            return []
        if isinstance(v, str):
            return [v]
        return v


class Experience(BaseModel):
    role: Optional[str] = None
    company: Optional[str] = None
    duration: Optional[str] = None
    description: Optional[List[str]] = Field(default_factory=list)

    @field_validator("description", mode="before")
    @classmethod
    def coerce_list(cls, v):
        if v is None:
            return []
        if isinstance(v, str):
            return [v]
        return v


class Club(BaseModel):
    name: Optional[str] = None
    role: Optional[str] = None
    description: Optional[str] = None


class Education(BaseModel):
    degree: Optional[str] = None
    institution: Optional[str] = None
    year: Optional[str] = None
    grade: Optional[str] = None


class Resume(BaseModel):
    name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    summary: Optional[str] = None
    skills: Optional[List[str]] = Field(default_factory=list)
    education: Optional[List[Education]] = Field(default_factory=list)
    projects: Optional[List[Project]] = Field(default_factory=list)
    experience: Optional[List[Experience]] = Field(default_factory=list)
    clubs: Optional[List[Club]] = Field(default_factory=list)

    @field_validator("skills", mode="before")
    @classmethod
    def coerce_skills(cls, v):
        if v is None:
            return []
        if isinstance(v, str):
            return [s.strip() for s in v.split(",") if s.strip()]
        return v

    @field_validator("education", "projects", "experience", "clubs", mode="before")
    @classmethod
    def coerce_list(cls, v):
        if v is None:
            return []
        return v


# ---------------------------------------------------------------------------
# 1. call_llm
# ---------------------------------------------------------------------------

def call_llm(prompt: str, system: str = "") -> str:
    """
    Send a prompt to OpenRouter and return the response text.
    Handles API errors gracefully, returns empty string on failure.
    """
    if not OPENROUTER_API_KEY:
        raise EnvironmentError(
            "OPENROUTER_API_KEY is not set. "
            "Export it as an environment variable before running."
        )

    messages = []
    if system:
        messages.append({"role": "system", "content": system})
    messages.append({"role": "user", "content": prompt})

    payload = {
        "model": DEFAULT_MODEL,
        "messages": messages,
        "temperature": 0.4,
        "max_tokens": 2048,
    }

    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://resume-builder.local",
        "X-Title": "AI Resume Builder",
    }

    try:
        response = requests.post(
            OPENROUTER_BASE_URL,
            headers=headers,
            json=payload,
            timeout=60,
        )
        response.raise_for_status()
        data = response.json()

        choices = data.get("choices")
        if not choices or not isinstance(choices, list):
            logger.error("Unexpected API response — no 'choices' field: %s", data)
            return ""

        content = choices[0].get("message", {}).get("content", "")
        return content.strip()

    except requests.exceptions.Timeout:
        logger.error("OpenRouter request timed out.")
        return ""
    except requests.exceptions.HTTPError as exc:
        logger.error("HTTP error from OpenRouter: %s", exc)
        return ""
    except requests.exceptions.RequestException as exc:
        logger.error("Network error: %s", exc)
        return ""
    except (KeyError, IndexError, ValueError) as exc:
        logger.error("Failed to parse OpenRouter response: %s", exc)
        return ""


# ---------------------------------------------------------------------------
# 2. get_next_question
# ---------------------------------------------------------------------------

# Mandatory fields collected in order
_MANDATORY_FIELDS = ["name", "email", "phone"]

# Optional resume fields (AI decides relevance)
_OPTIONAL_FIELDS = ["skills", "education", "projects", "experience", "clubs"]

# Minimum data thresholds before generating resume
_MIN_OPTIONAL_FIELDS_FILLED = 5


def _count_filled(data: dict, fields: list) -> int:
    """Count how many fields from the list are present and non-empty."""
    count = 0
    for f in fields:
        val = data.get(f)
        if val and (not isinstance(val, (list, dict)) or len(val) > 0):
            count += 1
    return count


def get_next_question(data: dict) -> str:
    """
    Dynamically decide the next question to ask the user.

    Enforces name → email → phone order first.
    Then uses the LLM to decide which resume section to ask about next.
    Returns "GENERATE_RESUME" when sufficient data is collected.
    """
    # --- Enforce mandatory fields in order ---
    for field in _MANDATORY_FIELDS:
        if not data.get(field):
            questions = {
                "name": "What is your full name?",
                "email": "What is your email address?",
                "phone": "What is your phone number?",
            }
            return questions[field]
        # FORCE CRITICAL FIELDS FIRST
    if not data.get("education"):
        return "What is your education? (degree, college, year)"

    if not data.get("projects"):
        return "Describe a project you have worked on."

    if not data.get("experience"):
        return "Do you have any work or internship experience?"

    if not data.get("clubs"):
        return "Are you part of any clubs or activities?"

    # --- Check if we already have enough to generate ---
    filled_optional = _count_filled(data, _OPTIONAL_FIELDS)
    if filled_optional >= _MIN_OPTIONAL_FIELDS_FILLED:
        return "GENERATE_RESUME"

    # --- Ask LLM for the next most relevant question ---
    collected = {k: v for k, v in data.items() if v}
    missing = [f for f in _OPTIONAL_FIELDS if not data.get(f)]

    if not missing:
        return "GENERATE_RESUME"

    system_prompt = (
    "You are a resume assistant. "
    "Ask ONE short, practical resume question. "
    "ONLY ask about skills, education, projects, experience, or clubs. "
    "NO theory, NO comparisons, NO explanations."
    )

    prompt = f"""
The user is building their resume.

Already collected:
{json.dumps(collected, indent=2)}

Still missing (choose ONE to ask about): {", ".join(missing)}

Rules:
- Ask only ONE question
- Keep it short and direct (under 15 words)
- Match the field to the user's context
- Do NOT ask philosophical or comparison questions
- Output ONLY the question text, nothing else

Question:
""".strip()

    question = call_llm(prompt, system=system_prompt)

    # Fallback: if LLM fails or gives garbage, ask about first missing field
    if not question or len(question) > 200 or "\n" in question:
        fallbacks = {
            "skills": "What are your key technical and soft skills?",
            "education": "What is your highest education? (degree, college, year)",
            "projects": "Describe a project you have built. (name, tech, what it does)",
            "experience": "Do you have any work experience? (role, company, duration)",
            "clubs": "Are you part of any clubs or student organizations?",
        }
        return fallbacks.get(missing[0], "Tell me more about your background.")

    # Strip leading/trailing quotes if LLM wraps the question in them
    question = question.strip('"').strip("'").strip()
    return question


# ---------------------------------------------------------------------------
# 3. generate_resume
# ---------------------------------------------------------------------------

def _clean_llm_json(raw: str) -> str:
    """
    Extract valid JSON from LLM output.
    Handles markdown fences, leading/trailing text, and BOM characters.
    """
    # Remove BOM if present
    raw = raw.lstrip("\ufeff").strip()

    # Strip markdown code fences
    raw = re.sub(r"^```(?:json)?", "", raw, flags=re.IGNORECASE).strip()
    raw = re.sub(r"```$", "", raw).strip()

    # Find the first { and last } to extract only the JSON object
    start = raw.find("{")
    end = raw.rfind("}")
    if start != -1 and end != -1 and end > start:
        return raw[start : end + 1]

    return raw


def _remove_empty_entries(resume_dict: dict) -> dict:
    """Remove list items that are essentially empty objects."""

    def is_empty_entry(entry: dict) -> bool:
        """Returns True if all values in the dict are None / empty."""
        for v in entry.values():
            if v and v not in ([], {}, ""):
                return False
        return True

    for key in ("projects", "experience", "clubs", "education"):
        if key in resume_dict and isinstance(resume_dict[key], list):
            resume_dict[key] = [
                item for item in resume_dict[key]
                if isinstance(item, dict) and not is_empty_entry(item)
            ]
    return resume_dict


def generate_resume(data: dict) -> dict:
    """
    Use the LLM to generate a structured, professional Resume object
    from the raw collected user data.
    """
    system_prompt = (
        "You are an expert resume writer. "
        "Output ONLY valid JSON. No explanation. No markdown. No extra text. "
        "Assume the user is a student unless stated otherwise."
    )

    prompt = f"""
You are given raw resume input from a user. Your job is to generate a polished, professional resume as STRICT JSON.

Raw user input:
{json.dumps(data, indent=2)}

Instructions:
1. Output ONLY a JSON object matching the schema below. No explanation, no markdown, no prose.
2. Expand weak or short inputs into professional, detailed resume language.
3. For every project and experience:
   - Write 3 bullet points
   - Each bullet MUST:
       • start with a strong action verb
       • include what was built/done
       • include impact or result (if possible)
4. Education must include multiple entries if possible (e.g., both college degree and high school).
5. Skills must be a flat list of strings.
6. If a field has no data, use an empty list [] or null.
7. Generate a 2-sentence professional summary for the "summary" field.

JSON Schema to follow exactly:
{{
  "name": "string",
  "email": "string",
  "phone": "string",
  "summary": "string",
  "skills": ["string"],
  "education": [
    {{
      "degree": "string",
      "institution": "string",
      "year": "string",
      "grade": "string or null"
    }}
  ],
  "projects": [
    {{
      "name": "string",
      "description": ["bullet string 1", "bullet string 2", "bullet string 3"],
      "technologies": ["string"],
      "link": "string or null"
    }}
  ],
  "experience": [
    {{
      "role": "string",
      "company": "string",
      "duration": "string",
      "description": ["bullet string 1", "bullet string 2"]
    }}
  ],
  "clubs": [
    {{
      "name": "string",
      "role": "string",
      "description": "string or null"
    }}
  ]
}}

Output ONLY valid JSON. No explanation.
""".strip()

    raw_output = call_llm(prompt, system=system_prompt)

    if not raw_output:
        logger.warning("LLM returned empty output. Building Resume from raw data.")
        return _resume_from_raw(data)

    cleaned = _clean_llm_json(raw_output)

    try:
        resume_dict = json.loads(cleaned)
    except json.JSONDecodeError as exc:
        logger.error("JSON decode error: %s\nRaw output snippet: %.300s", exc, raw_output)
        return _resume_from_raw(data)

    resume_dict = _remove_empty_entries(resume_dict)

    try:
        resume = Resume.model_validate(resume_dict)
    except Exception as exc:
        logger.error("Pydantic validation error: %s", exc)
        # Best-effort: merge LLM data with raw data
        resume = _resume_from_raw({**data, **resume_dict})

    return resume.model_dump()


def _resume_from_raw(data: dict) -> Resume:
    """Fallback: build a minimal Resume directly from raw collected data."""
    return Resume(
        name=data.get("name"),
        email=data.get("email"),
        phone=data.get("phone"),
        summary=None,
        skills=data.get("skills", []),
        education=[],
        projects=[],
        experience=[],
        clubs=[],
    )


# ---------------------------------------------------------------------------
# 4. export_to_docx — Professional Layout
# ---------------------------------------------------------------------------

def _set_font(run, name: str = "Calibri", size: int = 11,
              bold: bool = False, color: Optional[RGBColor] = None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def _add_horizontal_rule(doc: DocxDocument):
    """Add a thin bottom-border paragraph as a visual divider."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(4)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E74B5")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_section_heading(doc: DocxDocument, title: str):
    """Add a styled section heading with a divider line beneath it."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(title.upper())
    _set_font(run, size=11, bold=True, color=RGBColor(0x2E, 0x74, 0xB5))
    _add_horizontal_rule(doc)


def _add_bullet(doc: DocxDocument, text: str, indent_level: int = 0):
    """Add a bullet-point paragraph."""
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.left_indent = Inches(0.25 + indent_level * 0.25)
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    run = para.add_run(text)
    _set_font(run, size=10.5)


def export_to_docx(resume: Resume) -> str:
    """
    Build a professionally formatted DOCX resume.
    Returns the path to the saved file.
    """
    doc = DocxDocument()

    # Page margins (narrow for resume)
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # ── Name ──────────────────────────────────────────────────────────────
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(2)
    name_run = name_para.add_run(resume.name or "Your Name")
    _set_font(name_run, size=22, bold=True, color=RGBColor(31,56,100))

    # ── Contact line ──────────────────────────────────────────────────────
    contact_parts = [p for p in [resume.email, resume.phone] if p]
    contact_line = "  |  ".join(contact_parts)
    if contact_line:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.paragraph_format.space_after = Pt(4)
        contact_run = contact_para.add_run(contact_line)
        _set_font(contact_run, size=10, color=RGBColor(0x44, 0x44, 0x44))

    # ── Summary ───────────────────────────────────────────────────────────
    if resume.summary:
        _add_section_heading(doc, "Professional Summary")
        summary_para = doc.add_paragraph()
        summary_para.paragraph_format.space_after = Pt(4)
        run = summary_para.add_run(resume.summary)
        _set_font(run, size=10.5)

    # ── Skills ────────────────────────────────────────────────────────────
    if resume.skills:
        _add_section_heading(doc, "Skills")
        skills_para = doc.add_paragraph()
        skills_para.paragraph_format.space_after = Pt(4)
        run = skills_para.add_run("  •  ".join(resume.skills))
        _set_font(run, size=10.5)

    # ── Education ─────────────────────────────────────────────────────────
    if resume.education:
        _add_section_heading(doc, "Education")
        for edu in resume.education:
            if not edu.degree and not edu.institution:
                continue
            edu_para = doc.add_paragraph()
            edu_para.paragraph_format.space_before = Pt(4)
            edu_para.paragraph_format.space_after = Pt(1)

            degree_run = edu_para.add_run(edu.degree or "Degree")
            _set_font(degree_run, size=10.5, bold=True)

            if edu.institution:
                inst_run = edu_para.add_run(f"  —  {edu.institution}")
                _set_font(inst_run, size=10.5)

            if edu.year or edu.grade:
                detail_parts = [p for p in [edu.year, edu.grade] if p]
                detail_run = edu_para.add_run("  |  " + "  |  ".join(detail_parts))
                _set_font(detail_run, size=10, color=RGBColor(0x66, 0x66, 0x66))

    # ── Projects ──────────────────────────────────────────────────────────
    if resume.projects:
        _add_section_heading(doc, "Projects")
        for project in resume.projects:
            if not project.name:
                continue
            proj_para = doc.add_paragraph()
            proj_para.paragraph_format.space_before = Pt(5)
            proj_para.paragraph_format.space_after = Pt(1)

            name_run = proj_para.add_run(project.name)
            _set_font(name_run, size=10.5, bold=True)

            if project.technologies:
                tech_run = proj_para.add_run(
                    "  |  Tech: " + ", ".join(project.technologies)
                )
                _set_font(tech_run, size=10, color=RGBColor(0x44, 0x44, 0x44))

            if project.link:
                link_run = proj_para.add_run(f"  |  {project.link}")
                _set_font(link_run, size=9.5, color=RGBColor(0x1A, 0x6A, 0xC8))

            for bullet in (project.description or []):
                _add_bullet(doc, bullet)

    # ── Experience ────────────────────────────────────────────────────────
    if resume.experience:
        _add_section_heading(doc, "Experience")
        for exp in resume.experience:
            if not exp.role and not exp.company:
                continue
            exp_para = doc.add_paragraph()
            exp_para.paragraph_format.space_before = Pt(5)
            exp_para.paragraph_format.space_after = Pt(1)

            role_run = exp_para.add_run(exp.role or "Role")
            _set_font(role_run, size=10.5, bold=True)

            if exp.company:
                company_run = exp_para.add_run(f"  —  {exp.company}")
                _set_font(company_run, size=10.5)

            if exp.duration:
                dur_run = exp_para.add_run(f"  |  {exp.duration}")
                _set_font(dur_run, size=10, color=RGBColor(0x66, 0x66, 0x66))

            for bullet in (exp.description or []):
                _add_bullet(doc, bullet)

    # ── Clubs & Activities ────────────────────────────────────────────────
    if resume.clubs:
        _add_section_heading(doc, "Clubs & Activities")
        for club in resume.clubs:
            if not club.name:
                continue
            club_para = doc.add_paragraph()
            club_para.paragraph_format.space_before = Pt(4)
            club_para.paragraph_format.space_after = Pt(1)

            club_run = club_para.add_run(club.name)
            _set_font(club_run, size=10.5, bold=True)

            if club.role:
                role_run = club_para.add_run(f"  —  {club.role}")
                _set_font(role_run, size=10.5)

            if club.description:
                desc_para = doc.add_paragraph()
                desc_para.paragraph_format.left_indent = Inches(0.25)
                desc_para.paragraph_format.space_before = Pt(1)
                desc_para.paragraph_format.space_after = Pt(2)
                desc_run = desc_para.add_run(club.description)
                _set_font(desc_run, size=10.5)

    # ── Save ──────────────────────────────────────────────────────────────
    safe_name = re.sub(r"[^\w\s-]", "", resume.name or "resume").strip().replace(" ", "_")
    filename = f"{safe_name}_{uuid.uuid4().hex[:6]}.docx"
    filepath = OUTPUT_DIR / filename
    doc.save(str(filepath))
    logger.info("DOCX saved: %s", filepath)
    return str(filepath)


# ---------------------------------------------------------------------------
# 5. export_to_pdf
# ---------------------------------------------------------------------------
# def build_html_resume(resume: Resume) -> str:
#     return f"""
#     <html>
#     <head>
#         <style>
#             body {{
#                 font-family: Arial, sans-serif;
#                 margin: 40px auto;
#                 max-width: 800px;
#                 line-height: 1.6;
#             }}
#             h1 {{
#                 text-align: center;
#                 color: #1f3864;
#                 margin-bottom: 5px;
#             }}
#             .contact {{
#                 text-align: center;
#                 color: #555;
#                 margin-bottom: 20px;
#             }}
#             h2 {{
#                 border-bottom: 2px solid #2E74B5;
#                 padding-bottom: 5px;
#                 margin-top: 25px;
#             }}
#             ul {{
#                 margin-top: 5px;
#             }}
#         </style>
#     </head>
#     <body>

#     <h1>{resume.name}</h1>
#     <div class="contact">{resume.email} | {resume.phone}</div>

#     <h2>Summary</h2>
#     <p>{resume.summary or ""}</p>

#     <h2>Skills</h2>
#     <p>{" • ".join(resume.skills)}</p>

#     <h2>Education</h2>
#     {"".join([
#         f"<p><b>{e.degree}</b> — {e.institution} | {e.year} {('| ' + e.grade) if e.grade else ''}</p>"
#         for e in resume.education
#     ])}

#     <h2>Projects</h2>
#     {"".join([
#         f"<p><b>{p.name}</b></p><ul>" +
#         "".join([f"<li>{d}</li>" for d in p.description]) +
#         "</ul>"
#         for p in resume.projects
#     ])}

#     <h2>Experience</h2>
#     {"".join([
#         f"<p><b>{e.role}</b> — {e.company} | {e.duration}</p><ul>" +
#         "".join([f"<li>{d}</li>" for d in e.description]) +
#         "</ul>"
#         for e in resume.experience
#     ])}

#     <h2>Clubs</h2>
#     {"".join([
#         f"<p><b>{c.name}</b> — {c.role}</p><p>{c.description or ''}</p>"
#         for c in resume.clubs
#     ])}

#     </body>
#     </html>
#     """

# def export_to_pdf(resume: Resume) -> str:
#     """
#     Convert resume to PDF via docx2pdf.
#     Falls back to DOCX if conversion fails or docx2pdf is not installed.
#     """
#     docx_path = export_to_docx(resume)

#     try:
#         from docx2pdf import convert  # type: ignore

#         pdf_path = docx_path.replace(".docx", ".pdf")
#         convert(docx_path, pdf_path)
#         logger.info("PDF saved: %s", pdf_path)
#         return pdf_path

#     except ImportError:
#         logger.warning(
#             "docx2pdf is not installed. Install it with: pip install docx2pdf. "
#             "Returning DOCX path instead."
#         )
#         return docx_path
#     except Exception as exc:
#         logger.error("PDF conversion failed: %s. Returning DOCX path.", exc)
#         return docx_path



from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4

def export_to_pdf(resume: Resume) -> str:
    filepath = OUTPUT_DIR / f"{resume.name.replace(' ', '_')}.pdf"

    doc = SimpleDocTemplate(str(filepath), pagesize=A4)
    styles = getSampleStyleSheet()
    elements = []

    # Name
    elements.append(Paragraph(f"<b>{resume.name}</b>", styles["Title"]))
    elements.append(Spacer(1, 10))

    # Contact
    elements.append(Paragraph(f"{resume.email} | {resume.phone}", styles["Normal"]))
    elements.append(Spacer(1, 10))

    # Summary
    if resume.summary:
        elements.append(Paragraph("<b>Summary</b>", styles["Heading2"]))
        elements.append(Paragraph(resume.summary, styles["Normal"]))
        elements.append(Spacer(1, 10))

    # Skills
    if resume.skills:
        elements.append(Paragraph("<b>Skills</b>", styles["Heading2"]))
        elements.append(Paragraph(", ".join(resume.skills), styles["Normal"]))
        elements.append(Spacer(1, 10))

    # Projects
    if resume.projects:
        elements.append(Paragraph("<b>Projects</b>", styles["Heading2"]))
        for p in resume.projects:
            elements.append(Paragraph(f"<b>{p.name}</b>", styles["Normal"]))
            bullets = ListFlowable(
                [Paragraph(b, styles["Normal"]) for b in p.description],
                bulletType='bullet'
            )
            elements.append(bullets)
            elements.append(Spacer(1, 8))

    # Experience
    if resume.experience:
        elements.append(Paragraph("<b>Experience</b>", styles["Heading2"]))
        for e in resume.experience:
            elements.append(Paragraph(f"<b>{e.role} - {e.company}</b>", styles["Normal"]))
            bullets = ListFlowable(
                [Paragraph(b, styles["Normal"]) for b in e.description],
                bulletType='bullet'
            )
            elements.append(bullets)
            elements.append(Spacer(1, 8))

    doc.build(elements)
    return str(filepath)

# ---------------------------------------------------------------------------
# 6. export_resume — Unified Dispatcher
# ---------------------------------------------------------------------------

def export_resume(resume: Resume, format: str = "docx") -> str:
    """
    Export the resume to the requested format.

    Args:
        resume: Validated Resume Pydantic model.
        format: "docx" or "pdf"

    Returns:
        Absolute path to the exported file.
    """
    fmt = format.lower().strip()
    if fmt == "pdf":
        return export_to_pdf(resume)
    elif fmt == "docx":
        return export_to_docx(resume)
    else:
        logger.warning("Unknown format '%s'. Defaulting to DOCX.", format)
        return export_to_docx(resume)


# # ---------------------------------------------------------------------------
# # Demo / Smoke Test
# # ---------------------------------------------------------------------------


# if __name__ == "__main__":
#     # Sample collected data
#     sample_data = {
#         "name": "Arjun Menon",
#         "email": "arjun.menon@example.com",
#         "phone": "+91-9876543210",
#         "skills": "Python, React, Machine Learning, SQL, Git",
#         "education": "B.Tech in Computer Science at NIT Calicut (2022–2026)",
#         "projects": "Built a WiFi congestion prediction model using Random Forest",
#         "experience": "Intern at TechCorp, built REST APIs for 3 months",
#         "clubs": "Google Developer Student Club, Technical Lead",
#     }

#     print("=" * 60)
#     print("AI RESUME BUILDER — Demo Run")
#     print("=" * 60)

#     # 1. Simulate the question flow
#     print("\n[1] Simulating question flow...")
#     session: dict = {}
#     fields_order = ["name", "email", "phone", "skills", "education", "projects", "experience", "clubs"]
#     for field in fields_order:
#         q = get_next_question(session)
#         if q == "GENERATE_RESUME":
#             break
#         print(f"   Q: {q}")
#         session[field] = sample_data[field]
#         print(f"   A: {session[field]}\n")

#     # 2. Generate resume via LLM
#     print("[2] Generating structured resume via LLM...")
#     resume_obj = generate_resume(sample_data)
    
#     # Validate with Pydantic
#     resume_model = Resume.model_validate(resume_obj)
    
#     print(f"   Name     : {resume_model.name}")
#     print(f"   Skills   : {resume_model.skills[:5]}")
#     print(f"   Projects : {len(resume_model.projects)} entries")

#     # 3. Export to DOCX
#     print("\n[3] Exporting to DOCX...")
#     docx_path = export_resume(resume_model, format="docx")
#     print(f"   Saved: {docx_path}")

#     # 4. Export to PDF
#     print("\n[4] Exporting to PDF (requires docx2pdf and MS Word)...")
#     pdf_path = export_resume(resume_model, format="pdf")
#     print(f"   Saved: {pdf_path}")

#     print("\nDone.")